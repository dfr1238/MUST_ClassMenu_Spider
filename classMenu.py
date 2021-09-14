from PySimpleGUI.PySimpleGUI import WIN_CLOSED
import selenium
from selenium import webdriver
import PySimpleGUI as sg
from docx import Document
from docx.shared import  Cm
from selenium.webdriver.support.ui import Select
import sys, os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
import selenium.webdriver.support.ui as ui
import webbrowser
import winreg

class spider_Gui:

    def set_Input_Ready_Window():
        input_Ready_Layout =[
            [sg.Text('請到新開啟的網頁中登入')]
        ]
        return sg.Window('準備爬取資料',input_Ready_Layout,finalize=True)
    
    def set_finish_Window(spider):
        finsih_Window_Layout =[
            [sg.Text(f'###已抓取完該學生課表資料，如需下一筆請關閉 Word 後到網頁再次登入###\n學生姓名：{spider.std_name}\t學號：{spider.std_id}\n學年：{spider.std_year}\t學期：{spider.std_season}\n請到程式目錄尋找 Word 檔：{spider.std_id} - {spider.std_name} - {spider.std_year} - {spider.std_season} 課表.docx')]
        ]
        return sg.Window('完成爬取',finsih_Window_Layout,finalize=True,modal=True)

    def set_running_Window(spider):
        running_Window_Layout =[
            [sg.Text(f'已偵測到可抓取表格！\n學生姓名：{spider.std_name}\t學號：{spider.std_id}\n學年：{spider.std_year}\t學期：{spider.std_season}')]
        ]
        return sg.Window('已偵測到爬取資料！',running_Window_Layout,finalize=True)
    pass

class classMenu_Spider:
    doc=''
    table_Element =''
    table_TrList=''
    std_name=''
    std_id=''
    std_year=''
    std_season=''
    std_class=''
    driver=''
    url='https://sss.must.edu.tw/pars_new/pars_index.asp'

    def __init__(self) -> None:
        chrome_options = webdriver.ChromeOptions()
        #chrome_options.add_argument('--headless')
        chrome_options.add_argument('--disable-gpu')
        if __name__ == "__main__":

            if getattr(sys, 'frozen', False): 
                chrome_driver_path = os.path.join('.\chromedriver.exe')
                print(chrome_driver_path)
                self.driver = webdriver.Chrome(executable_path=chrome_driver_path,options=chrome_options)
            else:
                try:
                    self.driver = webdriver.Chrome(options=chrome_options)
                except selenium.common.exceptions.SessionNotCreatedException:
                    sg.popup_error('WebDriver 版本錯誤或未搜尋到，程式將下載對應 的 WebDriver。壓縮檔下載後，請將壓縮黨內的 chromedriver 放置與該程式同個目錄下',keep_on_top=True)
                    reg_key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,r'Software\Google\Chrome\BLBeacon')
                    ch_ver = winreg.QueryValueEx(reg_key,'version')[0]
                    webbrowser.open(f'https://chromedriver.storage.googleapis.com/{ch_ver}/chromedriver_win32.zip', new=2)
                    sys.exit()
                    
        try:
            self.driver.get(self.url)
            self.driver.maximize_window()
            self.driver.set_page_load_timeout(10)
        except selenium.common.exceptions.WebDriverException or selenium.common.TimeoutException:
            sg.popup_error(f'建立網頁驅動器時發生問題！請檢查網路連線與網頁 {self.url} 的狀態！')
            os._exit(0)
        wait = ui.WebDriverWait(self.driver,10)
        wait.until(lambda driver: driver.find_element_by_name('STDNO'))
        print(f'連線至 {self.url} 成功！')

    def creating_Word(self): #建立 Word 檔
        self.driver.minimize_window()
        self.doc = Document()
        style = self.doc.styles['Normal']
        font = style.font
        font.size = Pt(12)
        p = self.doc.add_paragraph(f'學號：{self.std_id}\t姓名：{self.std_name}\t班級：')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = self.doc.styles['Normal']
        p = self.doc.add_paragraph(f'學年：{self.std_year}\t學期：{self.std_season}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = self.doc.styles['Normal']
        table = self.doc.add_table(rows=1,cols=8)
        table.style = 'Light Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text ='時段'
        hdr_cells[1].text ='星期一'
        hdr_cells[2].text ='星期二'
        hdr_cells[3].text ='星期三'
        hdr_cells[4].text ='星期四'
        hdr_cells[5].text ='星期五'
        hdr_cells[6].text ='星期六'
        hdr_cells[7].text ='星期日'
        hdr_cells[0].width = Cm(.5)
        for i in range(1,8):
            hdr_cells[i].width = Cm(3)
        #date_heading=[['時段'],['星期一'],['星期二'],['星期三'],['星期四'],['星期五'],['星期六'],['星期日']]
        #print(date_heading)
        first_row=True
        for row in self.table_TrList:
            if(first_row):
                first_row=False
                continue
            tdlist = row.find_elements_by_tag_name('td')
            sg.one_line_progress_meter(f'匯出成 Word 檔中...',self.table_TrList.index(row),len(self.table_TrList)-1,'Progress',f'學年：{self.std_year}\t學期：{self.std_season}\n目前爬取學生名：{self.std_name}\t學號：{self.std_id}')
            row_cells = table.add_row().cells
            td_count=0
            face_night=False
            for td in tdlist:
                temp_text = td.text.split("\n")
                if(len(temp_text) > 3):
                    del temp_text[1:3]
                if(str("".join(temp_text)) == "進 修 部"):
                    print(str("".join(temp_text)))
                    face_night=True
                    p=row_cells[td_count].add_paragraph(str("\n".join(temp_text)))
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                else:
                    row_cells[td_count].text = str("\n".join(temp_text))
                td_count+=1
            if(face_night):
                row_cells[0].merge(row_cells[-1])
            #print('\n')
        sections = self.doc.sections
        for section in sections: #調整邊界
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
        self.doc.save(f'.\{self.std_id} - {self.std_name} - {self.std_year} - {self.std_season} 課表.docx')
        pass

    def check_table(self):
        try:
            wait = ui.WebDriverWait(self.driver,0.1)
            wait.until(lambda driver: driver.find_element_by_xpath('//*[@id="Windows-TOP"]/div/table'))
            self.table_Element = self.driver.find_element_by_xpath('//*[@id="Windows-TOP"]/div/table')
            self.table_TrList = self.driver.find_elements_by_tag_name('tr')
            info = self.driver.find_element_by_xpath('/html/body/div[1]/div/div[2]/ul/li/a/span/small').text.split("\n")
            year_List= Select(self.driver.find_element_by_name('CosYear'))
            season_List = Select(self.driver.find_element_by_name('CosSmtr'))
            self.std_year = year_List.first_selected_option.text
            self.std_season = season_List.first_selected_option.text
            self.std_name = info[1].replace('家長','')
            self.std_id = info[0]
            return True
        except selenium.common.exceptions.TimeoutException:
            print('尚未找到課表元素！')
            return False
        except selenium.common.exceptions.UnexpectedAlertPresentException:
            return False
            pass
        except selenium.common.exceptions.WebDriverException:
            sys.exit()

    def showing_data(self):
        sg.popup_notify(f'姓名：{self.std_name}\t學號：{self.std_id}',title='已找到課表！',display_duration_in_ms=150,fade_in_duration=150)
        pass

    def waiting_Input(self):
        ready_Window=None
        ready_Window=spider_Gui.set_Input_Ready_Window()
        running_Window=None
        finish_Window=None
        while True:
            window , event , values = sg.read_all_windows(500)
            state=self.check_table()
            if window == finish_Window:
                if event == WIN_CLOSED:
                    self.driver.quit()
                    window.close()
                    break
            if window == ready_Window:
                if event == WIN_CLOSED:
                    self.driver.quit()
                    window.close()
                    break
            if state:
                if(finish_Window!=None):
                    finish_Window.close()
                self.showing_data()
                ready_Window.close()
                running_Window=spider_Gui.set_running_Window(self)
                self.creating_Word()
                running_Window.close()
                self.driver.back()
                self.driver.maximize_window()
                docx=(f'.\{self.std_id} - {self.std_name} - {self.std_year} - {self.std_season} 課表.docx')
                os.startfile(docx)
                finish_Window=spider_Gui.set_finish_Window(self)

Spider = classMenu_Spider()
Spider.waiting_Input()